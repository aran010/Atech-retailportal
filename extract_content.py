#!/usr/bin/env python3
"""Extract full text content from 9 Word documents."""

import os
import traceback
from docx import Document

FILES = [
    "/Users/aran/Downloads/AC Receipt.docx",
    "/Users/aran/Downloads/AFFIDAVIT (prop).docx",
    "/Users/aran/Downloads/AFFIDAVIT(Regd. Pharmacist).docx",
    "/Users/aran/Downloads/Camera Receipt -.docx",
    "/Users/aran/Downloads/Inverter Receipt.docx",
    "/Users/aran/Downloads/PROP WORKING REPORT.docx",
    "/Users/aran/Downloads/Refrigerator Receipt.docx",
    "/Users/aran/Downloads/Rent Agreement.doc",
    "/Users/aran/Downloads/RP WORKING REPORT.docx",
]


def extract_docx(filepath):
    """Extract all content from a .docx file."""
    doc = Document(filepath)

    # --- Paragraphs (body) ---
    print("=== BODY PARAGRAPHS ===")
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "None"
        text = para.text
        print(f"  P{i} [{style}]: {text}")

    # --- Tables ---
    if doc.tables:
        print(f"\n=== TABLES ({len(doc.tables)} found) ===")
        for ti, table in enumerate(doc.tables):
            print(f"  --- Table {ti} ---")
            for ri, row in enumerate(table.rows):
                cells_text = [cell.text for cell in row.cells]
                print(f"    Row {ri}: {cells_text}")
    else:
        print("\n=== TABLES: None ===")

    # --- Headers & Footers ---
    print("\n=== HEADERS & FOOTERS ===")
    for si, section in enumerate(doc.sections):
        print(f"  --- Section {si} ---")
        # Header
        try:
            header = section.header
            if header and not header.is_linked_to_previous:
                for pi, para in enumerate(header.paragraphs):
                    print(f"    Header P{pi}: {para.text}")
                # tables inside header
                for ti, table in enumerate(header.tables):
                    print(f"    Header Table {ti}:")
                    for ri, row in enumerate(table.rows):
                        cells_text = [cell.text for cell in row.cells]
                        print(f"      Row {ri}: {cells_text}")
            else:
                print("    Header: (linked to previous or empty)")
        except Exception as e:
            print(f"    Header error: {e}")

        # Footer
        try:
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for pi, para in enumerate(footer.paragraphs):
                    print(f"    Footer P{pi}: {para.text}")
                for ti, table in enumerate(footer.tables):
                    print(f"    Footer Table {ti}:")
                    for ri, row in enumerate(table.rows):
                        cells_text = [cell.text for cell in row.cells]
                        print(f"      Row {ri}: {cells_text}")
            else:
                print("    Footer: (linked to previous or empty)")
        except Exception as e:
            print(f"    Footer error: {e}")


def main():
    for filepath in FILES:
        basename = os.path.basename(filepath)
        print("\n" + "#" * 80)
        print(f"# FILE: {basename}")
        print(f"# PATH: {filepath}")
        print("#" * 80)

        if not os.path.exists(filepath):
            print("  *** FILE NOT FOUND ***")
            continue

        try:
            extract_docx(filepath)
        except Exception as e:
            print(f"  *** ERROR processing file: {e} ***")
            traceback.print_exc()


if __name__ == "__main__":
    main()
