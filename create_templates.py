"""
create_templates.py
-------------------
Run this script ONCE to generate all 9 .docx templates inside the
`templates/` directory. Each file is pre-formatted with Jinja2-style tags
that docxtpl understands (e.g. {{r prop_name }}).

Usage:
    python create_templates.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")


def _set_default_font(doc):
    """Set default font to Calibri 11pt for the entire document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    # Also set the East Asia font
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:cs"), "Calibri")


def _add_title(doc, text):
    """Add a bold, centered title paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def _add_para(doc, text, bold=False, alignment=None):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if alignment:
        p.alignment = alignment
    return p


def _add_empty_para(doc):
    """Add an empty paragraph for spacing."""
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════
# 1. AC RECEIPT
# ═══════════════════════════════════════════════════════════════════════════
def _template_ac_receipt(doc):
    _add_title(doc, "AC RECEIPT")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r ac_seller_name }} {{r ac_seller_relation }} situated at "
        "{{r ac_seller_address }}, have received {{r ac_amount }} "
        "({{r ac_amount_words }} only) From {{r prop_name }} {{r prop_relation }} "
        "Sh. {{r prop_father_name }} R/o {{r prop_address }} through its Prop "
        "M/s {{r firm_name }} situated at {{r firm_address }} Against Sale of my "
        "AC, make {{r ac_make }} {{r ac_details }} in working condition which was "
        "purchased by me and received full and Final Payment in cash."
    )
    _add_empty_para(doc)
    _add_empty_para(doc)
    _add_para(doc, "(Sign of purchaser)\t\t\t\t\t(Sign of Seller)")
    _add_para(doc, "Date\t\t\t\t\t\t\tDate")


# ═══════════════════════════════════════════════════════════════════════════
# 2. AFFIDAVIT (Prop)
# ═══════════════════════════════════════════════════════════════════════════
def _template_affidavit_prop(doc):
    _add_title(doc, "AFFIDAVIT ({% if entity_val == 'Partner' %}Partner{% else %}Prop{% endif %})")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} "
        "R/o {{r prop_address }}, do hereby solemnly affirm and declare as under: -"
    )
    _add_empty_para(doc)

    # 15 numbered paragraphs
    _add_para(doc,
        "1. That I have never been convicted by any court in India under the "
        "Drugs & Cosmetics Act 1940 and Rules 1945 framed there under; -"
    )
    _add_para(doc,
        "2. That I am {% if entity_val == 'Partner' %}partner{% else %}sole proprietor{% endif %} of the firm M/s {{r firm_name }} Situated at "
        "{{r firm_address }}, and do hereby applying for grant of new retail sale Drug License."
    )
    _add_para(doc,
        "3. That I myself will be the overall in-charge and responsible person to my "
        "said firm for its day to day conduct and control of business of the firm."
    )
    _add_para(doc,
        "{% if property_ownership == 'Owned' %}"
        "4. That the sale premises of my said firm is the owned property and the same "
        "premise is under my legal possession/occupancy and it is not connected to any residence."
        "{% else %}"
        "4. That the sale premises of my said firm is the rented property and the same "
        "premise is under my legal possession/occupancy as a tenant and it is not "
        "connected to any residence."
        "{% endif %}"
    )
    _add_para(doc,
        "5. That I had never been a Prop or an active or sleeping partner at any such "
        "firm, whose Retail Sale/wholesale drugs License had ever been cancelled by the "
        "licensing authority for any reason whatsoever."
    )
    _add_para(doc,
        "6. That the firm has appointed Mr. {{r rp_name }} {{r rp_relation }} Sh. "
        "{{r rp_father_name }} R/o {{r rp_address }}, to sell by way of retail sale, "
        "who has joined as Regd. Pharmacist at a salary of Rs. {{r rp_salary }}/- per "
        "month w.e.f. {{r rp_joining_date }} and he is regd. Pharmacist from "
        "{{r rp_pharmacy_council }} vide Regn. No. {{r rp_reg_number }} dated "
        "{{r rp_reg_date }} which is valid up {{r rp_reg_valid_upto }}."
    )
    _add_para(doc,
        "7. That I shall pay the salary of the competent person by way of "
        "cheque/online transfer in his account."
    )
    _add_para(doc,
        "8. That I have installed CCTV camera at my shop/premises and I will keep "
        "one-month backup recording of CCTV Camera."
    )
    _add_para(doc,
        "9. That I have installed a refrigerator at my firm, which is in working "
        "condition and wooden racks in my said firm for the storage of Drugs."
    )
    _add_para(doc,
        "10. That I opt and want to keep all records of Sale of drugs in cash memos "
        "/bills/invoice of my said firm which shall be maintained properly and in "
        "legible manner."
    )
    _add_para(doc,
        "11. That the sale premises of my said firm will not be used/utilized for any "
        "other purpose expect for business of those categories of drugs which will "
        "include in the License applied for by me or granted to me at my said firm."
    )
    _add_para(doc,
        "12. That I shall comply with the provisions, rules regulation and conditions "
        "of the Drugs & Cosmetics Act 1940 and Rules 1945 framed there under for the "
        "time being in force or are amended from time to time under the said Act & Rules."
    )
    _add_para(doc,
        "13. That I shall obtain new Drug License in case of any change in constitution "
        "or premises takes place at my firm. I shall inform the Licensing Authorities "
        "if any area alterations take place at my firm."
    )
    _add_para(doc,
        "14. That if in case of resignation of competent person of my firm sale will "
        "not done in the absence of RP and I will appoint new RP immediately and will "
        "give written information to the drugs Dep't within one month."
    )
    _add_para(doc,
        "15. That if in case I close my firm I will give written information along with "
        "list of drugs laying at my firm unsold."
    )

    _add_empty_para(doc)
    _add_para(doc, "DEPONENT")
    _add_empty_para(doc)
    _add_para(doc,
        "Verification: - I the above named do hereby solemnly affirm and declare that "
        "whatsoever is started above is true and correct the best of my knowledge and "
        "belief and nothing has been concealed therein."
    )
    _add_empty_para(doc)
    _add_para(doc, "Date: -")
    _add_para(doc, "Place: -\t\t\t\t\tDEPONENT")


# ═══════════════════════════════════════════════════════════════════════════
# 3. AFFIDAVIT (Regd. Pharmacist)
# ═══════════════════════════════════════════════════════════════════════════
def _template_affidavit_rp(doc):
    _add_title(doc, "AFFIDAVIT (Regd. Pharmacist)")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r rp_name }} {{r rp_relation }} Sh. {{r rp_father_name }} "
        "R/o {{r rp_address }}, do hereby solemnly affirms and declare as under: -"
    )
    _add_empty_para(doc)

    # 13 numbered paragraphs
    _add_para(doc,
        "1. That I have never been convicted by any court in India under the "
        "Drugs & cosmetic Act 1940 and rules 1945 framed there under."
    )
    _add_para(doc,
        "2. That I have joined as registered with the firm M/s {{r firm_name }} "
        "Situated at {{r firm_address }}, as whole time bases on salary Rs. "
        "{{r rp_salary }}/-per month w.e.f. {{r rp_joining_date }}."
    )
    _add_para(doc,
        "3. That I had never been a proprietor or an active or sleeping partner at "
        "any such firm, whose, retail sale/whole sale drugs license had ever been "
        "cancelled by the licensing authority for any reason whatsoever."
    )
    _add_para(doc,
        "4. That I have passed {{r rp_qualification }} from {{r rp_college }} and I am "
        "registered from {{r rp_pharmacy_council }} Panchkula vide Regn. No. "
        "{{r rp_reg_number }} dated {{r rp_reg_date }} which is valid up "
        "{{r rp_reg_valid_upto }}."
    )
    _add_para(doc,
        "5. That I will renew my registration certificate from H.S.P.C five years "
        "and I will submit the proof to the licensing authority."
    )
    _add_para(doc,
        "6. That I will not work at any other firm/any institute/in any drug "
        "manufacturing unit in any capacity during my service with this firm."
    )
    _add_para(doc,
        "7. That I am not a student of any educational institute."
    )
    _add_para(doc,
        "{% if rp_prev_firm_name or rp_prev_firm_address or rp_resign_date %}"
        "8. That previously I was worked as regd. pharmacist with the firm M/s "
        "{{r rp_prev_firm_name }} situated at {{r rp_prev_firm_address }}, and resign "
        "from the firm on dated {{r rp_resign_date }}. After resign from this firm, "
        "I have not worked as Regd. Pharmacist at anywhere anyfirm."
        "{% else %}"
        "8. That I have not worked as Regd. pharmacist at any firm since my registration as Regd. pharmacist. "
        "Before Joining from this firm, I have not worked as Regd. Pharmacist at anywhere anyfirm"
        "{% endif %}"
    )
    _add_para(doc,
        "9. That all particulars of my said qualification and registration are true "
        "on the basis of documents and certificate possessed and submitted by me and "
        "the same are genuine and not bogus, fake or forged."
    )
    _add_para(doc,
        "10. That I shall comply-with the provision, rules, regulation and condition "
        "of the Drugs & Cosmetics act 1940 and Rules 1945 framed there under for the "
        "time being in force or are amended from time to time under the said acts and rules."
    )
    _add_para(doc,
        "11. That if in case I resign from the said firm, I will give written "
        "information to the licensing authority with the consent of prop of the "
        "firm one month before."
    )
    _add_para(doc,
        "12. That at present I am residing at {{r rp_address }}."
    )
    _add_para(doc,
        "13. That I will work at the firm as a full-time employee, attending "
        "regularly from my given address."
    )

    _add_empty_para(doc)
    _add_para(doc, "Deponent")
    _add_empty_para(doc)
    _add_para(doc,
        "Verification: I the above named do hereby solemnly affirm and declare to "
        "that whatever is started above is true and correct to the best of my "
        "knowledge and nothing has been concealed therein."
    )
    _add_empty_para(doc)
    _add_para(doc, "Date: -")
    _add_para(doc, "Place: -\t\t\t\t\tDeponent")


# ═══════════════════════════════════════════════════════════════════════════
# 4. Camera Receipt
# ═══════════════════════════════════════════════════════════════════════════
def _template_camera_receipt(doc):
    _add_title(doc, "CAMERA RECEIPT")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r cam_seller_name }} {{r cam_seller_relation }} R/o "
        "{{r cam_seller_address }}, have received {{r cam_amount }} (Rupees "
        "{{r cam_amount_words }} only) From {{r prop_name }} {{r prop_relation }} "
        "Sh. {{r prop_father_name }} R/o {{r prop_address }} through its Prop "
        "M/s {{r firm_name }} Situated at {{r firm_address }} against Sale of my "
        "Camera, make {{r cam_make }} in working condition which was purchased "
        "by me and received full and Final Payment in cash."
    )
    _add_empty_para(doc)
    _add_empty_para(doc)
    _add_para(doc, "(Sign of purchaser)\t\t\t\t\t(Sign of Seller)")
    _add_para(doc, "Date\t\t\t\t\t\t\tDate")


# ═══════════════════════════════════════════════════════════════════════════
# 5. Inverter Receipt
# ═══════════════════════════════════════════════════════════════════════════
def _template_inverter_receipt(doc):
    _add_title(doc, "INVERTER RECEIPT")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r inv_seller_name }} {{r inv_seller_relation }} R/o "
        "{{r inv_seller_address }}, have received {{r inv_amount }} "
        "({{r inv_amount_words }} only) From {{r prop_name }} {{r prop_relation }} "
        "Sh. {{r prop_father_name }} R/o {{r prop_address }} through its prop "
        "M/s {{r firm_name }} situated at {{r firm_address }} Against Sale of my "
        "Inverter, make {{r inv_make }} in working condition which was purchased "
        "by me and received full and Final Payment in cash."
    )
    _add_empty_para(doc)
    _add_empty_para(doc)
    _add_para(doc, "(Sign of purchaser)\t\t\t\t\t(Sign of Seller)")
    _add_para(doc, "Date\t\t\t\t\t\t\tDate")


# ═══════════════════════════════════════════════════════════════════════════
# 6. PROP WORKING REPORT
# ═══════════════════════════════════════════════════════════════════════════
def _template_prop_working_report(doc):
    _add_para(doc, "To whom soever it may concern")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} "
        "R/o {{r prop_address }}, do hereby certify that during the last three "
        "years my working details are as follows:"
    )
    _add_empty_para(doc)

    # Create table: header row + loop-start row + data row + loop-end row
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    # Row 0: Header row
    headers = ["Sr. No.", "Time Period", "Working/Occupation", "Remarks"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # Row 1: Loop start tag (must be alone in its own row)
    table.rows[1].cells[0].paragraphs[0].text = "{%tr for item in prop_work_history %}"
    table.rows[1].cells[1].paragraphs[0].text = ""
    table.rows[1].cells[2].paragraphs[0].text = ""
    table.rows[1].cells[3].paragraphs[0].text = ""

    # Row 2: Data row with item values
    data_row = table.rows[2]
    data_row.cells[0].text = ""
    data_row.cells[0].paragraphs[0].add_run(
        "{{r item.sr_no }}"
    ).font.size = Pt(11)
    data_row.cells[1].text = ""
    data_row.cells[1].paragraphs[0].add_run(
        "{{r item.time_period }}"
    ).font.size = Pt(11)
    data_row.cells[2].text = ""
    data_row.cells[2].paragraphs[0].add_run(
        "{{r item.occupation }}"
    ).font.size = Pt(11)
    data_row.cells[3].text = ""
    data_row.cells[3].paragraphs[0].add_run(
        "{{r item.remarks }}"
    ).font.size = Pt(11)

    # Row 3: Loop end tag (must be alone in its own row)
    table.rows[3].cells[0].paragraphs[0].text = "{%tr endfor %}"
    table.rows[3].cells[1].paragraphs[0].text = ""
    table.rows[3].cells[2].paragraphs[0].text = ""
    table.rows[3].cells[3].paragraphs[0].text = ""

    _add_empty_para(doc)
    _add_para(doc, "Phone no-{{r prop_phone }}")
    _add_empty_para(doc)
    _add_para(doc, "Date: -……………\t\t\t\t\tProprietor Signature")
    _add_empty_para(doc)
    _add_para(doc, "Verification")
    _add_empty_para(doc)
    _add_para(doc,
        "I Know Mr. {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} "
        "having personally for the last three years and as per best of my knowledge, "
        "the above-mentioned details are correct and he has never been "
        "prosecuted/convicted by any Court in India."
    )
    _add_empty_para(doc)

    # Witness 1 (left) and Witness 2 (right) — use a 2-column table for alignment
    witness_table = doc.add_table(rows=4, cols=2)
    witness_labels = ["Name……………………………", "Add.……………………………", "Designation……………………………", "Date……………………………"]
    for i, label in enumerate(witness_labels):
        witness_table.rows[i].cells[0].text = label
        witness_table.rows[i].cells[1].text = label
        for cell in [witness_table.rows[i].cells[0], witness_table.rows[i].cells[1]]:
            for run in cell.paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════════════════
# 7. Refrigerator Receipt
# ═══════════════════════════════════════════════════════════════════════════
def _template_fridge_receipt(doc):
    _add_title(doc, "REFRIGERATOR RECEIPT")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r fridge_seller_name }} {{r fridge_seller_relation }} R/o "
        "{{r fridge_seller_address }}, have received {{r fridge_amount }} (Rupees "
        "{{r fridge_amount_words }} only) From {{r prop_name }} {{r prop_relation }} "
        "Sh. {{r prop_father_name }} R/o {{r prop_address }} through its Prop "
        "M/s {{r firm_name }} Situated at {{r firm_address }} against Sale of my "
        "Fridge make {{r fridge_make }} {{r fridge_details }} in working condition "
        "which was purchased by me and received full and Final Payment in cash."
    )
    _add_empty_para(doc)
    _add_empty_para(doc)
    _add_para(doc, "(Sign of purchaser)\t\t\t\t\t(Sign of Seller)")
    _add_para(doc, "Date\t\t\t\t\t\t\tDate")


# ═══════════════════════════════════════════════════════════════════════════
# 8. Rent Agreement
# ═══════════════════════════════════════════════════════════════════════════
def _template_rent_agreement(doc):
    _add_title(doc, "RENT AGREEMENT")
    _add_empty_para(doc)

    _add_para(doc,
        "This Rent Agreement is made on this {{r rent_agreement_date }} "
        "(date of rent Agreement) by {{r landlord_name }} {{r landlord_relation }} "
        "Sh. {{r landlord_relative_name }} R/o {{r landlord_address }}. Herein after "
        "called the Lessor / Owner, Party of the first part"
    )
    _add_empty_para(doc)
    _add_para(doc, "AND", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_empty_para(doc)
    _add_para(doc,
        "{{r firm_name }} through its Prop {{r prop_name }} {{r prop_relation }} "
        "Sh. {{r prop_father_name }} R/o {{r prop_address }}. Herein after called "
        "Lessee/Tenant, Party of the Second Part."
    )
    _add_empty_para(doc)
    _add_para(doc,
        "That the expression of the term, Lessor/Owner and the Lessee/Tenant Shall "
        "mean and include their legal heir's successors, assigns, representative etc. "
        "whereas the Lessor/Owner is the owner and in possession of the property "
        "{{r shop_address }} and has agreed to let out the one shop at ground floor on "
        "said property, to the Lessee/Tenant and the Lessee/Tenant has agreed to take "
        "the same on rent of Rs. {{r rent_amount }}/- (Rupees {{r rent_amount_words }} "
        "only) per month."
    )
    _add_empty_para(doc)
    _add_para(doc, "That the Exact Location of this firm: -")
    _add_para(doc, "Right\t\t{{r neighbor_right }}")
    _add_para(doc, "Left\t\t{{r neighbor_left }}")
    _add_para(doc, "Front\t\t{{r neighbor_front }}")
    _add_para(doc, "Back\t\t{{r neighbor_back }}")
    _add_empty_para(doc)

    _add_para(doc, "NOW THIS RENT AGREEMENT WITNESSETH AS UNDER: -", bold=True)
    _add_empty_para(doc)

    _add_para(doc,
        "1. That the Lessor/Owner has agreed to let out the said Shop/premises "
        "to the Lessee/Tenant for the purpose of running a Retail Drug Business "
        "and the Lessee/Tenant has agreed to take the same on rent of Rs. "
        "{{r rent_amount }}/- (Rupees {{r rent_amount_words }} only) per month."
    )
    _add_para(doc,
        "2. That this Rent Agreement shall be for a period of {{r lease_months_words }} "
        "({{r lease_months }}) Months commencing from {{r rent_start_date }}."
    )
    _add_para(doc,
        "3. That the Lessee/Tenant shall pay the rent to the Lessor/Owner on or "
        "before the 10th of every month without any default."
    )
    _add_para(doc,
        "4. That the Lessee/Tenant shall use the premises only for the purpose "
        "mentioned above and shall not use it for any other purpose without the "
        "prior written consent of the Lessor/Owner."
    )
    _add_para(doc,
        "5. That the Lessee/Tenant shall not sub-let or transfer the tenancy "
        "rights to any third party without the prior written consent of the "
        "Lessor/Owner."
    )
    _add_para(doc,
        "6. That the Lessee/Tenant shall maintain the premises in good condition "
        "and shall hand over the same in the same condition at the time of vacating "
        "the premises subject to normal wear and tear."
    )
    _add_para(doc,
        "7. That the Lessee/Tenant shall not make any structural changes or "
        "alterations in the premises without the prior written consent of the "
        "Lessor/Owner."
    )
    _add_para(doc,
        "8. That the Lessor/Owner shall be responsible for all major repairs and "
        "the Lessee/Tenant shall be responsible for minor repairs and maintenance."
    )
    _add_para(doc,
        "9. That the electricity and water charges shall be borne by the "
        "Lessee/Tenant."
    )
    _add_para(doc,
        "10. That the Lessee/Tenant shall vacate the premises after the expiry of "
        "the lease period or on giving one month's prior notice in writing."
    )
    _add_para(doc,
        "11. That in case of any dispute arising out of this agreement, the same "
        "shall be resolved amicably and if not resolved, it shall be subject to "
        "the jurisdiction of the local courts."
    )
    _add_para(doc,
        "12. That both the parties have read and understood the terms and conditions "
        "of this agreement and have signed the same in the presence of the witnesses."
    )

    _add_empty_para(doc)
    _add_para(doc,
        "IN WITNESS WHEREOF the parties hereto have set their respective hands and "
        "signatures on this {{r rent_agreement_date }} in the presence of the "
        "following witnesses."
    )

    _add_empty_para(doc)
    _add_empty_para(doc)

    # Signature block using a 2-column table for alignment
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.rows[0].cells[0].text = "{{r landlord_name }}"
    sig_table.rows[0].cells[1].text = "{{r prop_name }}"
    sig_table.rows[1].cells[0].text = "(Name of the landlord)"
    sig_table.rows[1].cells[1].text = "(Name of the Tenant)"
    sig_table.rows[2].cells[0].text = "Lessor"
    sig_table.rows[2].cells[1].text = "Lessee"
    for row in sig_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)

    _add_empty_para(doc)
    _add_para(doc, "Witnesses:", bold=True)
    _add_para(doc, "1. ________________________")
    _add_para(doc, "2. ________________________")


# ═══════════════════════════════════════════════════════════════════════════
# 9. RP WORKING REPORT
# ═══════════════════════════════════════════════════════════════════════════
def _template_rp_working_report(doc):
    _add_para(doc, "To whom so ever it may concern")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r rp_name }} {{r rp_relation }} Sh. {{r rp_father_name }} "
        "R/o {{r rp_address }}, do hereby certify that during the last three "
        "years my working details are as follows:"
    )
    _add_empty_para(doc)

    # Create table: header row + loop-start row + data row + loop-end row
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    # Row 0: Header row
    headers = ["Sr. No.", "Time Period", "Working/Occupation", "Remarks"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)

    # Row 1: Loop start tag (must be alone in its own row)
    table.rows[1].cells[0].paragraphs[0].text = "{%tr for item in rp_work_history %}"
    table.rows[1].cells[1].paragraphs[0].text = ""
    table.rows[1].cells[2].paragraphs[0].text = ""
    table.rows[1].cells[3].paragraphs[0].text = ""

    # Row 2: Data row with item values
    data_row = table.rows[2]
    data_row.cells[0].text = ""
    data_row.cells[0].paragraphs[0].add_run(
        "{{r item.sr_no }}"
    ).font.size = Pt(11)
    data_row.cells[1].text = ""
    data_row.cells[1].paragraphs[0].add_run(
        "{{r item.time_period }}"
    ).font.size = Pt(11)
    data_row.cells[2].text = ""
    data_row.cells[2].paragraphs[0].add_run(
        "{{r item.occupation }}"
    ).font.size = Pt(11)
    data_row.cells[3].text = ""
    data_row.cells[3].paragraphs[0].add_run(
        "{{r item.remarks }}"
    ).font.size = Pt(11)

    # Row 3: Loop end tag (must be alone in its own row)
    table.rows[3].cells[0].paragraphs[0].text = "{%tr endfor %}"
    table.rows[3].cells[1].paragraphs[0].text = ""
    table.rows[3].cells[2].paragraphs[0].text = ""
    table.rows[3].cells[3].paragraphs[0].text = ""

    _add_empty_para(doc)
    _add_para(doc,
        "It is further certified that at present I have joined M/s {{r firm_name }} "
        "Situated at {{r firm_address }}, w.e.f {{r rp_joining_date }} as R.P. and "
        "neither working nor studying anywhere else."
    )
    _add_empty_para(doc)
    _add_para(doc, "Phone no-{{r rp_phone }}")
    _add_empty_para(doc)
    _add_para(doc, "Date: ……………………………\t\t\t\t\t(Regd. Pharmacist)")
    _add_empty_para(doc)
    _add_para(doc, "Verification")
    _add_empty_para(doc)
    _add_para(doc,
        "I know Mr. {{r rp_name }} {{r rp_relation }} Sh. {{r rp_father_name }} "
        "having qualification {{r rp_qualification }} personally for the last more "
        "than three years and as per best of my knowledge, the above-mentioned "
        "details are correct and he has never been prosecuted/convicted by any "
        "Court in India."
    )
    _add_empty_para(doc)

    # Witness section — 2-column table for alignment
    witness_table = doc.add_table(rows=4, cols=2)
    witness_labels = ["Name……………………………", "Add.……………………………", "Designation……………………………", "Date……………………………"]
    for i, label in enumerate(witness_labels):
        witness_table.rows[i].cells[0].text = label
        witness_table.rows[i].cells[1].text = label
        for cell in [witness_table.rows[i].cells[0], witness_table.rows[i].cells[1]]:
            for run in cell.paragraphs[0].runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════════════════
# ADDRESS CHANGE — AFFIDAVIT (Prop)
# ═══════════════════════════════════════════════════════════════════════════
def _template_addr_affidavit_prop(doc):
    _add_title(doc, "AFFIDAVIT ({% if entity_val == 'Partner' %}Partner{% else %}Prop{% endif %})")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} "
        "R/o {{r prop_address }}, do hereby solemnly affirm and declare as under: -"
    )
    _add_empty_para(doc)

    _add_para(doc,
        "That I have never been convicted by any court in India under the "
        "Drugs & Cosmetics Act 1940 and Rules 1945 framed there under; -"
    )
    _add_para(doc,
        "That I am {% if entity_val == 'Partner' %}partner{% else %}sole proprietor{% endif %} of the firm following details:-"
    )
    _add_empty_para(doc)

    # ── Firm details table (5 cols) ──────────────────────────────────────
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    headers = ["Firm Name", "Old Address", "Drug license Number and Validity",
               "Applying for", "New Address"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    # Data row
    data = table.rows[1]
    fields = ["{{r firm_name }}", "{{r old_address }}",
              "{{r drug_license_number }}", "{{r applying_for }}",
              "{{r new_address }}"]
    for i, f in enumerate(fields):
        data.cells[i].text = ""
        data.cells[i].paragraphs[0].add_run(f).font.size = Pt(11)

    _add_empty_para(doc)
    _add_para(doc,
        "That I myself will be the overall in-charge and responsible person to my "
        "said firm for its day to day conduct and control of business of the firm."
    )
    _add_para(doc,
        "{% if property_ownership == 'Owned' %}"
        "That the sale premises of my said firm is the owned property and the same "
        "premise is under my legal possession/occupancy and it is not connected to any residence."
        "{% else %}"
        "That the sale premises of my said firm is the rented property and the same "
        "premise is under my legal possession/occupancy as tenant and it is not "
        "connected to any residence."
        "{% endif %}"
    )
    _add_para(doc,
        "That I had never been a Prop./Partner/director/auth. Signatory at any such "
        "firm, whose Retail Sale/wholesale drugs License had ever been cancelled by "
        "the licensing authority for any reason whatsoever."
    )
    _add_para(doc,
        "That the firm has already appointed is the following Regd. pharmacist: -"
    )
    _add_empty_para(doc)

    # ── RP details table (5 cols) ────────────────────────────────────────
    rp_table = doc.add_table(rows=2, cols=5)
    rp_table.style = "Table Grid"
    rp_headers = ["RP/CP Name", "Address", "Regn. no. and date", "Validity", "Salary"]
    for i, h in enumerate(rp_headers):
        cell = rp_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    rp_data = rp_table.rows[1]
    rp_fields = [
        "{{r rp_name }} {{r rp_relation }} Sh. {{r rp_father_name }}",
        "{{r rp_address }}",
        "{{r rp_reg_number }}\n{{r rp_reg_date }}",
        "{{r rp_reg_valid_upto }}",
        "{{r rp_salary }}/-",
    ]
    for i, f in enumerate(rp_fields):
        rp_data.cells[i].text = ""
        rp_data.cells[i].paragraphs[0].add_run(f).font.size = Pt(11)

    _add_empty_para(doc)
    _add_para(doc,
        "That I shall pay the salary of the competent person by way of "
        "cheque/online transfer in his account."
    )
    _add_para(doc,
        "That I have installed CCTV camera at my shop/premises and I will keep "
        "one-month backup recording of CCTV Camera."
    )
    _add_para(doc,
        "That I have installed a refrigerator at my firm, which is in working "
        "condition and wooden racks in my said firm for the storage of Drugs."
    )
    _add_para(doc,
        "That I opt and want to keep all records of Sale of drugs in cash memos "
        "/bills/invoice of my said firm which shall be maintained properly and in "
        "legible manner."
    )
    _add_para(doc,
        "That the sale premises of my said firm will not be used/utilized for any "
        "other purpose expect for business of those categories of drugs which will "
        "include in the License applied for by me or granted to me at my said firm."
    )
    _add_para(doc,
        "That I shall comply with the provisions, rules regulation and conditions "
        "of the Drugs & Cosmetics Act 1940 and Rules 1945 framed there under for the "
        "time being in force or are amended from time to time under the said Act & Rules."
    )
    _add_para(doc,
        "That I shall obtain new Drug License in case of any change in constitution "
        "or premises takes place at my firm. I shall inform the Licensing Authorities "
        "if any area alterations take place at my firm."
    )
    _add_para(doc,
        "That if in case of resignation of competent person of my firm sale will "
        "not done in the absence of RP and I will appoint new RP immediately and will "
        "give written information to the drugs Dep't within one month."
    )
    _add_para(doc,
        "That if in case I close my firm I will give written information along with "
        "list of drugs laying at my firm unsold."
    )

    _add_empty_para(doc)
    _add_para(doc, "\t\t\t\t\t\t\t\t\t\t\t\t\tDEPONENT")
    _add_para(doc, "Verification: -")
    _add_para(doc,
        "I the above named do hereby solemnly affirm and declare that whatsoever is "
        "started above is true and correct the best of my knowledge and belief and "
        "nothing has been concealed therein."
    )
    _add_empty_para(doc)
    _add_para(doc, "\t\t\t\tDate: -")
    _add_para(doc, "\t\t\t\tPlace: -\t\t\t\t\t\t\t\t\tDEPONENT")


# ═══════════════════════════════════════════════════════════════════════════
# ADDRESS CHANGE — AFFIDAVIT (Regd. Pharmacist)
# ═══════════════════════════════════════════════════════════════════════════
def _template_addr_affidavit_rp(doc):
    _add_title(doc, "AFFIDAVIT (Regd.Pharmacist)")
    _add_empty_para(doc)
    _add_para(doc,
        "I, {{r rp_name }} {{r rp_relation }} Sh. {{r rp_father_name }} "
        "R/o {{r rp_address }}, do hereby solemnly affirms and declare as under:"
    )
    _add_empty_para(doc)

    _add_para(doc,
        "That I have never been convicted by any court in India under the "
        "Drugs & cosmetic Act 1940 and rules 1945 framed there under."
    )
    _add_para(doc,
        "That I am working as registered Pharmacist with the firm following details: -"
    )
    _add_empty_para(doc)

    # ── Firm details table (6 cols) ──────────────────────────────────────
    table = doc.add_table(rows=2, cols=6)
    table.style = "Table Grid"
    headers = ["Firm Name", "Old Address", "Drug license Number and Validity",
               "Applying for", "New Address", "Salary"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    data = table.rows[1]
    fields = ["{{r firm_name }}", "{{r old_address }}",
              "{{r drug_license_number }}", "{{r applying_for }}",
              "{{r new_address }}", "{{r rp_salary }}/-"]
    for i, f in enumerate(fields):
        data.cells[i].text = ""
        data.cells[i].paragraphs[0].add_run(f).font.size = Pt(11)

    _add_empty_para(doc)
    _add_para(doc,
        "That I had never been a proprietor/partner/competent person at any such "
        "firm, whose, retail sale/whole sale drugs license had ever been cancelled "
        "by the licensing authority for any reason whatsoever."
    )
    _add_para(doc,
        "That my qualification is following details: -"
    )
    _add_empty_para(doc)

    # ── Qualification table (3 cols) ─────────────────────────────────────
    qual_table = doc.add_table(rows=2, cols=3)
    qual_table.style = "Table Grid"
    qual_headers = ["College name", "Registration Number Registration Date", "Validity"]
    for i, h in enumerate(qual_headers):
        cell = qual_table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    qual_data = qual_table.rows[1]
    qual_fields = [
        "{{r rp_college }}",
        "{{r rp_reg_number }}\n{{r rp_reg_date }}",
        "{{r rp_reg_valid_upto }}",
    ]
    for i, f in enumerate(qual_fields):
        qual_data.cells[i].text = ""
        qual_data.cells[i].paragraphs[0].add_run(f).font.size = Pt(11)

    _add_empty_para(doc)
    _add_para(doc,
        "That I will renew my registration certificate from H.S.P.C every five "
        "years and I will submit the proof to the Licensing authority."
    )
    _add_para(doc,
        "That I will not work at any other firm/any institute/in any drug "
        "manufacturing unit in any capacity during my service with this firm."
    )
    _add_para(doc,
        "That I am not a student of any educational institute."
    )
    _add_para(doc,
        "That presently I am working as regd. Pharmacist with the same firm "
        "at old premises."
    )
    _add_para(doc,
        "That all particulars of my said qualification and registration are true "
        "on the basis of documents and certificate possessed and submitted by me "
        "and the same are genuine and not bogus, fake or forged."
    )
    _add_para(doc,
        "That I shall comply-with the provision, rules, regulation and condition "
        "of the Drugs & Cosmetics act 1940 and Rules 1945 framed there under for "
        "the time being in force or are amended from time to time under the said "
        "acts and rules."
    )
    _add_para(doc,
        "That if in case I resign from the said firm, I will give written "
        "intimation to the licensing authority with the consent of prop of the "
        "firm one month before."
    )

    _add_empty_para(doc)
    _add_para(doc, "\tDeponent")
    _add_para(doc, "Verification: -")
    _add_para(doc,
        "I the above named do hereby solemnly affirm and declare to that whatever "
        "is started above is true and correct to the best of my knowledge and "
        "nothing has been concealed therein."
    )
    _add_empty_para(doc)
    _add_para(doc, "Date: -")
    _add_para(doc, "Place: -\t\t\t\t\t\t\t\t\t\tDeponent")


def _template_affidavit_director(doc: docx.Document):
    _add_para(doc, "AFFIDAVIT (DIRECTOR)", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} R/o {{r prop_address }}, "
        "do hereby solemnly affirm & declare as under:"
    )
    _add_para(doc,
        "That neither I nor any other director have ever been convicted by any court in India under the Drug & Cosmetics Act 1940 and Rules 1945 framed there under."
    )
    _add_para(doc,
        "That I am director of the firm M/s {{r firm_name }} situated at {{r firm_address }}, and do hereby applying for grant of new retail sale drug Licenses."
    )
    _add_para(doc, "That the other directors of the firm are the following detail:-")
    
    _add_para(doc,
        "{% for dir in directors_data %}"
        "{{r dir.name }} {{r dir.relation }} Sh. {{r dir.father_name }} R/o {{r dir.address }}."
        "{% endfor %}"
    )
    
    _add_para(doc,
        "That I myself{% if auth_signatory %} and {{r auth_signatory.name }}{% endif %} will be the overall in-charge and responsible person to our said firm for its day to day conduct & Control of business."
    )
    
    _add_para(doc,
        "{% if property_ownership == 'Owned' %}"
        "That the sale premises of my said firm is the owned property and the same premise is under my legal possession/occupancy and it is not connected to any residence."
        "{% else %}"
        "That the sale premises of my said firm is the rented property and the same premise is under my legal possession/occupancy and it is not connected to any residence."
        "{% endif %}"
    )

    _add_para(doc,
        "That I had never been a proprietor or an active or sleeping partner at any such firm, whose wholesale/retail sale drugs license had ever been cancelled by the licensing authority for any reason whatsoever."
    )
    _add_para(doc, "That the firm has appointed are the following Regd. Pharmacists:-")
    
    _add_para(doc,
        "{% for rp in pharmacists_data %}"
        "{{r rp.name }} {{r rp.relation }} Sh. {{r rp.father_name }} R/o {{r rp.address }}, to sell by way of retail sale, who has joined my firm as regd. Pharmacist at a salary of Rs. {{r rp.salary }}/- per month w.e.f {{r rp.joining_date }} and {% if rp.relation == 'D/o' or rp.relation == 'W/o' %}she{% else %}he{% endif %} is registered on H.S.P.C vide Regn. No. {{r rp.reg_no }} dated {{r rp.reg_date }} which is valid up to dated {{r rp.reg_valid_upto }}."
        "{% endfor %}"
    )
    
    _add_para(doc,
        "That I shall pay the salary of the competent person by way of cheque/online transfer in his account."
    )
    _add_para(doc,
        "That I have installed CCTV camera at my shop/premises and I will keep one-month backup recording of CCTV Camera."
    )
    _add_para(doc,
        "That the sale premises of my said firm will not be used/utilized for any other purpose expect for business of those categories of drugs which will include in the License applied for by me or granted to me at my said firm."
    )
    _add_para(doc,
        "That I shall comply with the provisions, rules regulation and conditions of the Drugs & Cosmetics Act 1940 and Rules1945 framed there under for the time being in force or are amended from time to time under the said Act & Rules."
    )
    _add_para(doc,
        "That I shall obtain new Drug License in case of any change in constitution or premises takes place at my firm. I shall inform the Licensing Authorities if any area alterations take place at my firm."
    )
    _add_para(doc,
        "That if in case of resignation of regd. Pharmacist of my firm sale will not done in the absence of regd. Pharmacist and I will appoint new regd. Pharmacist immediately and will give written information to the drugs Dep’t within one month."
    )
    _add_para(doc,
        "That if in case I close my firm I will give written information along with list of drugs laying at my firm unsold."
    )
    
    p_dep = doc.add_paragraph()
    p_dep.add_run("DEPONENT").bold = True
    p_dep.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _add_para(doc, "Verification:-")
    _add_para(doc,
        "I the above named do hereby solemnly affirm and declare that whatsoever is started above is true and correct the best of my knowledge and belief and nothing has been concealed therein."
    )
    
    _add_para(doc, "Date:-\nPlace: -")
    p_dep2 = doc.add_paragraph()
    p_dep2.add_run("DEPONENT").bold = True
    p_dep2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _template_affidavit_auth(doc: docx.Document):
    _add_para(doc, "AFFIDAVIT (Auth. Signatory)", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} R/o {{r prop_address }}, "
        "do hereby solemnly affirm & declare as under:"
    )
    _add_para(doc,
        "That neither I nor any directors have ever been convicted by any court in India under the Drug & Cosmetics Act 1940 and Rules 1945 framed there under."
    )
    _add_para(doc,
        "That I am auth. Signatory of the firm M/s {{r firm_name }} situated at {{r firm_address }}, and do hereby applying for grant of new retail sale drug Licenses."
    )
    _add_para(doc, "That the directors of the firm are the following detail:-")
    
    _add_para(doc,
        "{% for dir in directors_data %}"
        "{{r dir.name }} {{r dir.relation }} Sh. {{r dir.father_name }} R/o {{r dir.address }}."
        "{% endfor %}"
    )
    
    _add_para(doc,
        "That I myself and {% if directors_data %}{{r directors_data[0].name }}{% else %}the Director{% endif %} will be the overall in-charge and responsible person to our said firm for its day to day conduct & Control of business."
    )
    
    _add_para(doc,
        "{% if property_ownership == 'Owned' %}"
        "That the sale premises of my said firm is the owned property and the same premise is under my legal possession/occupancy and it is not connected to any residence."
        "{% else %}"
        "That the sale premises of my said firm is the rented property and the same premise is under my legal possession/occupancy and it is not connected to any residence."
        "{% endif %}"
    )

    _add_para(doc,
        "That I had never been a proprietor or an active or sleeping partner at any such firm, whose wholesale/retail sale drugs license had ever been cancelled by the licensing authority for any reason whatsoever."
    )
    _add_para(doc, "That the firm has appointed are the following Regd. Pharmacists:-")
    
    _add_para(doc,
        "{% for rp in pharmacists_data %}"
        "{{r rp.name }} {{r rp.relation }} Sh. {{r rp.father_name }} R/o {{r rp.address }}, to sell by way of retail sale, who has joined my firm as regd. Pharmacist at a salary of Rs. {{r rp.salary }}/- per month w.e.f {{r rp.joining_date }} and {% if rp.relation == 'D/o' or rp.relation == 'W/o' %}she{% else %}he{% endif %} is registered on H.S.P.C vide Regn. No. {{r rp.reg_no }} dated {{r rp.reg_date }} which is valid up to dated {{r rp.reg_valid_upto }}."
        "{% endfor %}"
    )
    
    _add_para(doc,
        "That I shall pay the salary of the competent person by way of cheque/online transfer in his account."
    )
    _add_para(doc,
        "That I have installed CCTV camera at my shop/premises and I will keep one-month backup recording of CCTV Camera."
    )
    _add_para(doc,
        "That the sale premises of my said firm will not be used/utilized for any other purpose expect for business of those categories of drugs which will include in the License applied for by me or granted to me at my said firm."
    )
    _add_para(doc,
        "That I shall comply with the provisions, rules regulation and conditions of the Drugs & Cosmetics Act 1940 and Rules1945 framed there under for the time being in force or are amended from time to time under the said Act & Rules."
    )
    _add_para(doc,
        "That I shall obtain new Drug License in case of any change in constitution or premises takes place at my firm. I shall inform the Licensing Authorities if any area alterations take place at my firm."
    )
    _add_para(doc,
        "That if in case of resignation of regd. Pharmacist of my firm sale will not done in the absence of regd. Pharmacist and I will appoint new regd. Pharmacist immediately and will give written information to the drugs Dep’t within one month."
    )
    _add_para(doc,
        "That if in case I close my firm I will give written information along with list of drugs laying at my firm unsold."
    )
    
    p_dep = doc.add_paragraph()
    p_dep.add_run("DEPONENT").bold = True
    p_dep.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _add_para(doc, "Verification:-")
    _add_para(doc,
        "I the above named do hereby solemnly affirm and declare that whatsoever is started above is true and correct the best of my knowledge and belief and nothing has been concealed therein."
    )
    
    _add_para(doc, "Date:-\nPlace: -")
    p_dep2 = doc.add_paragraph()
    p_dep2.add_run("DEPONENT").bold = True
    p_dep2.alignment = WD_ALIGN_PARAGRAPH.RIGHT



def _template_affidavit_partner(doc: docx.Document):
    _add_para(doc, "AFFIDAVIT (PARTNER)", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} R/o {{r prop_address }}, do hereby solemnly affirm & declare as under: -"
    )
    _add_para(doc,
        "That I have never been convicted by any court in India under the Drug & Cosmetics Act 1940 and Rules 1945 framed there under."
    )
    _add_para(doc,
        "That I am active partner of the firm M/s {{r firm_name }} Situated at {{r firm_address }}, and hereby applying for grant of new retail sale drug license. "
    )
    _add_para(doc, "That the other partner of the firm are the following detail:-")
    _add_para(doc,
        "{% for p in partners_data %}"
        "Mr./Ms. {{r p.name }} {{r p.relation }} Sh. {{r p.father_name }} R/o {{r p.address }}."
        "{% endfor %}"
    )
    _add_para(doc,
        "That I myself and {% if partners_data %}{{r partners_data[0].name }}{% else %}the Partner{% endif %} will be the overall in-charge and responsible person to our said firm for its day to day conduct and control of business."
    )
    _add_para(doc,
        "{% if property_ownership == 'Owned' %}"
        "That the sale premises of my said firm is the owned property and the same premises is under my legal possession/occupancy and it is not connected to any residence."
        "{% else %}"
        "That the sale premises of my said firm is the rented property and the same premises is under my legal possession/occupancy as a tenant and it is not connected to any residence."
        "{% endif %}"
    )
    _add_para(doc,
        "That I had never been a partner or an active or sleeping partner at any such firm, whose whole sale/retail sale Drug license had ever been cancelled by the licensing authority for any reason, whatsoever."
    )
    _add_para(doc, "That the firm has appointed are the following Regd. Pharmacists:-")
    _add_para(doc,
        "{% for rp in pharmacists_data %}"
        "{{r rp.name }} {{r rp.relation }} Sh. {{r rp.father_name }} R/o {{r rp.address }}, to sell by way of retail sale, who has joined my firm as regd. Pharmacist and {% if rp.relation == 'D/o' or rp.relation == 'W/o' %}she{% else %}he{% endif %} is regd. Pharmacist from H.S.P.C vide Regn. No. {{r rp.reg_no }} dated {{r rp.reg_date }} which is valid up {{r rp.reg_valid_upto }}."
        "{% endfor %}"
    )
    _add_para(doc,
        "That I would keep all records of sale & purchase etc. of drugs in cash memos / bills / invoices of my said firm, which shall be maintained properly and in legible manner."
    )
    _add_para(doc,
        "That the sale premises of my said firm will not be used / utilized for any other purpose expect for business of those categories of drugs which include in the license applied for by me or granted to me at my said firm."
    )
    _add_para(doc,
        "That I shall comply with the provision, rules, regulation and condition of the Drugs & Cosmetics Act 1940 and Rules 1945 framed there under for the time being in force or are amended from time to time under the said Acts and Rules."
    )
    _add_para(doc,
        "That I shall obtain new Drug License in case of any change in constitution or premises takes place at my firm. I shall inform the Licensing Authority, if any area alteration takes place at my firm."
    )
    _add_para(doc,
        "That the sale will not be done in the absence of Regd. Pharmacist in case of his resignation and I will appoint new Regd. Pharmacist immediately under written Intimation to the Licensing authority, within one month of such change."
    )
    _add_para(doc,
        "That if in case, I close my firm I will give written information along with list lying at my firm unsold."
    )
    p_dep = doc.add_paragraph()
    p_dep.add_run("Deponent").bold = True
    p_dep.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _add_para(doc, "Verification:")
    _add_para(doc,
        "I, the above named do hereby solemnly affirm and declare that whatever is started above is true and correct to the best of my knowledge and belief and nothing cancelled therein."
    )
    _add_para(doc, "Place: \nDated: ")
    p_dep2 = doc.add_paragraph()
    p_dep2.add_run("Deponent").bold = True
    p_dep2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _template_partner_working_report(doc: docx.Document):
    _add_para(doc, "To whom soever it may concern", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc,
        "I, {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} R/o {{r prop_address }}, do hereby certify that during the last three years my working details are as follows:  "
    )
    _add_para(doc, "Phone no-({{r prop_phone }})")
    
    p = doc.add_paragraph("Date: -……………						")
    p.add_run("(Partner Signature)").bold = True

    _add_para(doc, "Verification", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc,
        "I Know, Mr./Ms. {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} having personally for the last three years and as per best of my knowledge, the above-mentioned details are correct and he/she has never been prosecuted/convicted by any Court in India."
    )
    _add_para(doc, "Name: - ………………………………………        Name: -..............................................................")
    _add_para(doc, "Add. /Designation: - ……………………….….     Add. /Designation: -……………….………….")
    _add_para(doc, "Date: - ………………………………………….     Date: -..........................................................")

def _template_partnership_deed(doc: docx.Document):
    _add_para(doc, "PARTNERSHIP DEED", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_para(doc, "This deed of partnership is made by between the following parties: -")
    
    _add_para(doc,
        "Mr./Ms. {{r prop_name }} {{r prop_relation }} Sh. {{r prop_father_name }} R/o {{r prop_address }}. (Hereinafter called the party of the first part). Active Partner"
    )
    _add_para(doc,
        "{% for p in partners_data %}"
        "Mr./Ms. {{r p.name }} {{r p.relation }} Sh. {{r p.father_name }} R/o {{r p.address }}. (Hereinafter called the party of the second part). Active Partner"
        "{% endfor %}"
    )
    
    _add_para(doc,
        "WHEREAS both the above parties have been decided to carrying on the Business of Pharmaceutical, veterinary and sanitary preparation, dietetic substances adapted for medical use, food for babies. Plaster, Materials for dressing materials for stoping teeth, dental wax, disinfectants preparation for destroying vermin, fungicide’s, herbicide under the name and style of M/s {{r firm_name }} on the following terms and conditions: -"
    )
    _add_para(doc, "NOW THIS DEED OF PARTNERSHIP WITNESSETH AS FOLLOWS: -")
    _add_para(doc,
        "That the business of partnership shall be carried on under the name and styles of M/s {{r firm_name }} Situated at {{r firm_address }}."
    )
    _add_para(doc,
        "The partners with their mutual consent and determination may also start business under any other name and style and also at any other place as mutually decided by the partner."
    )
    _add_para(doc,
        "That the business of the partnership firm shall be of carrying on the Business of pharmaceutical, veterinary and sanitary preparation, dietetic substances adapted for medical use, food for babies. Plaster, Materials for dressing materials for stoping teeth, dental wax, disinfectants preparation for destroying vermin, fungicides, herbicide. The partners with their mutual consent and determination may also start any other business which they may deem profitable."
    )
    _add_para(doc,
        "That the business of the partnership firm shall be deemed to have commenced w.e.f as shall {{r partnership_start_date }}."
    )
    _add_para(doc, "That the partnership firm shall be at will.")
    _add_para(doc, "That the both partner contribute the capital of the Partnership firm agreed between them.")
    _add_para(doc,
        "The net profit of the partnership firm as per the account maintained by the firm after deducting of all expenses relating to trading activities or business of the partnership including rent. Salaries and other establishment expenses as well as interest and remuneration payable to the partners in accordance with the deed of partnership, year in the following proportion: -"
    )
    _add_para(doc, "The loss, if any, including loss of capital suffered in any year shall also is apportioned in the above proportion.")
    _add_para(doc,
        "Necessary capital as well as further funds required for the purpose of the partnership business shall be contributed or arranged by the partners in such manner as may be mutually agreed upon by and between the partners from the time to time. Interest at the rate of 12% annum or as may be prescribed under section 40 (B) (iv) of the income tax Act, 1961 or any other applicable provision as may be in force in income Tax assessment of the partners on the amount standing to the credit of the amount of the partners such interest shall be calculated and credited to the account of each partners at the close of the accounting year."
    )
    _add_para(doc,
        "All the parties of the partnership have agreed to keep themselves actively engaged in conduction the affairs of the business of the partnership firm, as working partners. It is hereby agreed that in consideration of all the parties working in the partnership, will be entitled to remuneration."
    )
    _add_para(doc,
        "The remuneration payable to the above said working partner shall be computed in the manner laid down in explanation 3 to section 40 (B) of the income tax act, 1961 or any other applicable provision as may be in force in the income tax assessment of the partnership firm for the relevant accounting year. Such remuneration shall be calculated at the close of the accounting year and shall be credited to the account of each working partner. The working partner shall be entitled to withdraw out of remuneration for their personal needs from time to time. The partners shall be entitled to increase or reduce the above remuneration and may agree to pay remuneration to other working partners or partners as the case may be. The parties hereto may also agree to revise the mode of calculating the above said remuneration as may be agreed to by the between the partners from time to time."
    )
    _add_para(doc, "That the accounting year of the firm shall be 31st March of every year.")
    _add_para(doc, "That once the balance sheet and profit and loss account prepared, checked and signed by all the partners and the same cannot be questioned by any of the partners thereafter.")
    _add_para(doc, "That the partnership firms open an account with any bank as the both partners here to may mutually agree shall be operated by.")
    _add_para(doc, "That the partners may be mutual consent introduce any other person in to partnership on such terms and conditions the existing partners may decide.")
    _add_para(doc,
        "That it will be open to any partner to retire from the partnership at any time after giving three months’ notice in writing, but will be liable to the partnership, according to his share for the firms, liabilities, if any as on the date of his expiry of his notice."
    )
    _add_para(doc,
        "That the death, insolvency, or retirement of the partner will not have the effect of disclosing the partnership so far as other partners are concerned. In the case of death, his legal heirs may if they so close, come in as partners subject to the terms and conditions thereof, by a notice, in writing given within three months of the death of the firm, in absence of such notice it will be taken that the legal heirs do not want to join the partnership in place of the deceased partner."
    )
    _add_para(doc, "That in the matters not hereby specifically agreed to the partnership will be governed by the provisions of the Indian Partnership Act 1932.")
    _add_para(doc,
        "In the event of any dispute between the partners with regard to anything Related to the partnership business, the matter should be refereed to Arbitration under the provision of the Indian Arbitration Act, 1932. "
    )
    _add_para(doc, "In WITNESS WHEREOF the parties here into have set their respective hands and seals, to these presents the day, month and year above written.")
    _add_para(doc, "WITNESSES")
    
    _add_para(doc,
        "\t\t\t\t\t\t{{r prop_name }}\n"
        "\t\t\t\t\t\t(Party of the First Part)\n"
        "{% for p in partners_data %}"
        "\t\t\t\t\t\t{{r p.name }}\n"
        "\t\t\t\t\t\t(Party of the Second Part)\n"
        "{% endfor %}"
    )


# ── Registry: filename → builder function ───────────────────────────────

TEMPLATES = {
    "AC Receipt.docx":                      _template_ac_receipt,
    "AFFIDAVIT (prop).docx":                _template_affidavit_prop,
    "AFFIDAVIT (Partner).docx":             _template_affidavit_partner,
    "Partner WORKING REPORT.docx":          _template_partner_working_report,
    "Partnership Deed.docx":                _template_partnership_deed,
    "AFFIDAVIT (Director).docx":            _template_affidavit_director,
    "AFFIDAVIT (Auth Signatory).docx":      _template_affidavit_auth,
    "AFFIDAVIT(Regd. Pharmacist).docx":     _template_affidavit_rp,
    "Camera Receipt -.docx":                _template_camera_receipt,
    "Inverter Receipt.docx":                _template_inverter_receipt,
    "PROP WORKING REPORT.docx":             _template_prop_working_report,
    "Refrigerator Receipt.docx":            _template_fridge_receipt,
    "Rent Agreement.docx":                  _template_rent_agreement,
    "RP WORKING REPORT.docx":               _template_rp_working_report,
}

# Address-change templates: replace the two affidavits, keep the rest
ADDR_CHANGE_TEMPLATES = {
    "AC Receipt.docx":                      _template_ac_receipt,
    "AFFIDAVIT (prop).docx":                _template_addr_affidavit_prop,
    "AFFIDAVIT (Partner).docx":             _template_affidavit_partner,
    "Partner WORKING REPORT.docx":          _template_partner_working_report,
    "Partnership Deed.docx":                _template_partnership_deed,
    "AFFIDAVIT (Director).docx":            _template_affidavit_director,
    "AFFIDAVIT (Auth Signatory).docx":      _template_affidavit_auth,
    "AFFIDAVIT(Regd. Pharmacist).docx":     _template_addr_affidavit_rp,
    "Camera Receipt -.docx":                _template_camera_receipt,
    "Inverter Receipt.docx":                _template_inverter_receipt,
    "PROP WORKING REPORT.docx":             _template_prop_working_report,
    "Refrigerator Receipt.docx":            _template_fridge_receipt,
    "Rent Agreement.docx":                  _template_rent_agreement,
    "RP WORKING REPORT.docx":               _template_rp_working_report,
}

ADDR_TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates_address_change")



def create_all_templates():
    # New File templates
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    for filename, builder_fn in TEMPLATES.items():
        doc = Document()
        _set_default_font(doc)
        builder_fn(doc)
        filepath = os.path.join(TEMPLATES_DIR, filename)
        doc.save(filepath)
        print(f"  ✓  Created: {filepath}")
    print(f"\nAll {len(TEMPLATES)} New File templates created in '{TEMPLATES_DIR}/'")

    # Address Change templates
    os.makedirs(ADDR_TEMPLATES_DIR, exist_ok=True)
    for filename, builder_fn in ADDR_CHANGE_TEMPLATES.items():
        doc = Document()
        _set_default_font(doc)
        builder_fn(doc)
        filepath = os.path.join(ADDR_TEMPLATES_DIR, filename)
        doc.save(filepath)
        print(f"  ✓  Created: {filepath}")
    print(f"\nAll {len(ADDR_CHANGE_TEMPLATES)} Address Change templates created in '{ADDR_TEMPLATES_DIR}/'")


if __name__ == "__main__":
    create_all_templates()

